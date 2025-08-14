import io
import os
import json
from datetime import datetime
from typing import List

import streamlit as st
from PIL import Image
import fitz  # PyMuPDF
import pytesseract
from deep_translator import GoogleTranslator
from langdetect import detect, DetectorFactory
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm

# Stabilizacja detekcji języka
DetectorFactory.seed = 0

APP_TITLE = "PDF OCR Studio (Streamlit) — by netdark_1966"
PAGE_ICON = "assets/icon.png"
LOGO_PATH = "assets/logo.png"
HISTORY_FILE = "history/conversions.json"
OUTPUT_DIR = "output"

# Ustaw ścieżkę do Tesseract na Windows:
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ------------- Utils -------------
def ensure_dirs():
    os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

def save_history(entry: dict):
    ensure_dirs()
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = []
    data.append(entry)
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def load_history() -> List[dict]:
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def clear_history():
    ensure_dirs()
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)

def render_page_thumb(pdf_doc: fitz.Document, page_index: int, dpi=110) -> Image.Image:
    page = pdf_doc.load_page(page_index)
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

def page_to_image(pdf_doc: fitz.Document, page_index: int, dpi=300) -> Image.Image:
    page = pdf_doc.load_page(page_index)
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

def detect_lang(text: str) -> str:
    try:
        return detect(text)
    except Exception:
        return "unknown"

def translate_text(text: str, source="auto", target="pl") -> str:
    try:
        return GoogleTranslator(source=source, target=target).translate(text)
    except Exception:
        # Tryb „best-effort”: jeśli tłumaczenie nie jest dostępne (offline/błąd), zwróć oryginał
        return text

def export_txt(pages_text: List[str], base_name: str) -> str:
    out_path = os.path.join(OUTPUT_DIR, f"{base_name}.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        for i, t in enumerate(pages_text, start=1):
            f.write(f"--- Strona {i} ---\n")
            f.write(t.rstrip() + "\n\n")
    return out_path

def export_docx(pages_text: List[str], base_name: str, footer_logo: str = None) -> str:
    doc = Document()
    doc.add_heading('PDF OCR Studio — wynik', level=1)
    for i, t in enumerate(pages_text, start=1):
        doc.add_heading(f"Strona {i}", level=2)
        for para in t.split("\n"):
            doc.add_paragraph(para)

    # Prosta sekcja „stopka/branding” na końcu dokumentu
    doc.add_page_break()
    doc.add_paragraph("Wygenerowano przez PDF OCR Studio — by netdark_1966")
    if footer_logo and os.path.exists(footer_logo):
        try:
            doc.add_picture(footer_logo, width=Inches(1.2))
        except Exception:
            pass

    out_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
    doc.save(out_path)
    return out_path

def export_pdf(pages_text: List[str], base_name: str) -> str:
    out_path = os.path.join(OUTPUT_DIR, f"{base_name}.pdf")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    left = 18 * mm
    top = height - 18 * mm
    max_width = width - 36 * mm
    line_height = 5.0 * mm

    for i, page in enumerate(pages_text, start=1):
        c.setFont("Times-Roman", 12)
        y = top
        c.drawString(left, y, f"Strona {i}")
        y -= line_height * 1.6

        for line in page.split("\n"):
            # prosty zawijacz linii
            content = line
            while content:
                fit = content
                while c.stringWidth(fit, "Times-Roman", 12) > max_width and len(fit) > 1:
                    fit = fit[:-1]
                c.drawString(left, y, fit)
                y -= line_height
                content = content[len(fit):]
                if y < 20 * mm:
                    c.showPage()
                    c.setFont("Times-Roman", 12)
                    y = top
            if y < 20 * mm:
                c.showPage()
                c.setFont("Times-Roman", 12)
                y = top
            else:
                y -= 1.0 * mm

        c.showPage()
    c.save()
    return out_path

def bytes_from_file(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()

def get_pdf_metadata(pdf_doc: fitz.Document) -> dict:
    """Wyciąga metadane z dokumentu PDF"""
    metadata = pdf_doc.metadata
    return {
        "title": metadata.get("title", "Brak"),
        "author": metadata.get("author", "Brak"),
        "subject": metadata.get("subject", "Brak"),
        "creator": metadata.get("creator", "Brak"),
        "producer": metadata.get("producer", "Brak"),
        "creation_date": metadata.get("creationDate", "Brak"),
        "modification_date": metadata.get("modDate", "Brak"),
        "pages": pdf_doc.page_count,
        "encrypted": pdf_doc.needs_pass
    }

# ------------- UI -------------
st.set_page_config(page_title=APP_TITLE, page_icon=PAGE_ICON, layout="wide")

# Nagłówek z brandingiem
col_logo, col_title = st.columns([1, 5])
with col_logo:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, use_container_width=True)
with col_title:
    st.title("📄 PDF OCR Studio")
    st.caption("by netdark_1966 — OCR, tłumaczenie EN→PL, eksport TXT/DOCX/PDF, historia, miniatury, dark mode")

with st.expander("ℹ️ Informacje i wskazówki", expanded=False):
    st.markdown(
        "- Wgraj PDF po lewej (sidebar).\n"
        "- Zaznacz strony do przetworzenia i ustaw opcje OCR/tłumaczenia.\n"
        "- Kliknij „🚀 Uruchom OCR” i obserwuj pasek postępu.\n"
        "- W sekcji „💾 Pobieranie wyników” pobierz wygenerowane pliki.\n"
        "- Zainstaluj Tesseract OCR i upewnij się, że jest w PATH."
    )

# Sidebar — sterowanie i pomoc
st.sidebar.header("⚙️ Ustawienia")
uploaded_pdf = st.sidebar.file_uploader("📥 Wgraj plik PDF", type=["pdf"])

ocr_lang_mode = st.sidebar.selectbox(
    "🈶 Język OCR",
    ["Auto (eng+pol, wykryj później)", "Angielski (eng)", "Polski (pol)", "Eng+Pol (eng+pol)"],
    index=0
)

translate_to_pl = st.sidebar.checkbox("🌐 Tłumacz wynik na polski (EN→PL)", value=True)

st.sidebar.markdown("---")
st.sidebar.subheader("📤 Eksport")

export_docx_opt = st.sidebar.checkbox("DOCX", value=True)
export_txt_opt = st.sidebar.checkbox("TXT", value=False)
export_pdf_opt = st.sidebar.checkbox("PDF", value=False)

# Inicjalizacja zmiennych
selected_pages: List[int] = []
pdf_doc = None

# Wczytanie PDF i wyświetlenie metadanych
if uploaded_pdf:
    pdf_bytes = uploaded_pdf.read()
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")

st.sidebar.markdown("---")
st.sidebar.subheader("📋 Metadane PDF")
if uploaded_pdf and pdf_doc:
    metadata = get_pdf_metadata(pdf_doc)
    st.sidebar.write(f"**Tytuł:** {metadata['title']}")
    st.sidebar.write(f"**Autor:** {metadata['author']}")
    st.sidebar.write(f"**Temat:** {metadata['subject']}")
    st.sidebar.write(f"**Twórca:** {metadata['creator']}")
    st.sidebar.write(f"**Producent:** {metadata['producer']}")
    st.sidebar.write(f"**Data utworzenia:** {metadata['creation_date']}")
    st.sidebar.write(f"**Data modyfikacji:** {metadata['modification_date']}")
    st.sidebar.write(f"**Liczba stron:** {metadata['pages']}")
    st.sidebar.write(f"**Zaszyfrowany:** {'Tak' if metadata['encrypted'] else 'Nie'}")
else:
    st.sidebar.caption("Wgraj plik PDF, aby zobaczyć metadane")

# Główna część — podgląd i selekcja stron
if uploaded_pdf and pdf_doc:
    st.success(f"✅ Wczytano: {uploaded_pdf.name} — {len(pdf_doc)} stron")

    st.subheader("📂 Podgląd stron")
    thumbs_cols = st.columns(4)
    selected_pages = []
    for idx in range(len(pdf_doc)):
        img = render_page_thumb(pdf_doc, idx, dpi=110)
        col = thumbs_cols[idx % 4]
        with col:
            st.image(img, caption=f"Strona {idx+1}", use_container_width=True)
            if st.checkbox(f"Wybierz {idx+1}", key=f"pick_{idx}"):
                selected_pages.append(idx)

    if not selected_pages:
        st.info("ℹ️ Nie wybrano stron — domyślnie przetworzymy wszystkie.")
        selected_pages = list(range(len(pdf_doc)))

    st.markdown("---")
    st.subheader("🛠️ Przetwarzanie")

    # Pasek postępu i status
    progress = st.progress(0)
    status = st.empty()

    # Start OCR
    if st.button("🚀 Uruchom OCR"):
        ensure_dirs()

        # Ustal język OCR
        if ocr_lang_mode == "Angielski (eng)":
            tess_lang = "eng"
        elif ocr_lang_mode == "Polski (pol)":
            tess_lang = "pol"
        elif ocr_lang_mode == "Eng+Pol (eng+pol)":
            tess_lang = "eng+pol"
        else:
            tess_lang = "eng+pol"  # auto-próba

        pages_text: List[str] = []
        sample_text = []

        # Plan kroków → do paska postępu
        steps_ocr = len(selected_pages)
        steps_translate = len(selected_pages) if translate_to_pl else 0
        steps_export = (1 if export_txt_opt else 0) + (1 if export_docx_opt else 0) + (1 if export_pdf_opt else 0)
        total_steps = max(1, steps_ocr + steps_translate + steps_export)
        done = 0

        # OCR
        status.info("🔎 OCR — trwa rozpoznawanie tekstu…")
        for i, pidx in enumerate(selected_pages, start=1):
            img = page_to_image(pdf_doc, pidx, dpi=300)
            text = pytesseract.image_to_string(img, lang=tess_lang)
            if i <= 3 and text.strip():
                sample_text.append(text.strip())
            pages_text.append(text)
            done += 1
            progress.progress(min(1.0, done / total_steps))

        # Detekcja języka (informacyjnie)
        detected = "unknown"
        if sample_text:
            detected = detect_lang("\n".join(sample_text))
        st.write(f"🧠 Rozpoznany język (próbka): {detected}")

        # Tłumaczenie
        do_translate = translate_to_pl and (detected in ["en", "unknown"] or ocr_lang_mode.startswith("Angielski"))
        if do_translate:
            status.info("🌐 Tłumaczenie EN→PL…")
            translated_pages = []
            for page_text in pages_text:
                translated_pages.append(translate_text(page_text, source="auto", target="pl"))
                done += 1
                progress.progress(min(1.0, done / total_steps))
            pages_text = translated_pages

        # Eksporty
        base_name = os.path.splitext(os.path.basename(uploaded_pdf.name))[0] + "_OCR"
        downloads = {}

        if export_txt_opt:
            status.info("💾 Eksport TXT…")
            txt_path = export_txt(pages_text, base_name)
            downloads["TXT"] = txt_path
            done += 1
            progress.progress(min(1.0, done / total_steps))

        if export_docx_opt:
            status.info("💾 Eksport DOCX…")
            docx_path = export_docx(pages_text, base_name, footer_logo=LOGO_PATH)
            downloads["DOCX"] = docx_path
            done += 1
            progress.progress(min(1.0, done / total_steps))

        if export_pdf_opt:
            status.info("💾 Eksport PDF…")
            pdf_path = export_pdf(pages_text, base_name)
            downloads["PDF"] = pdf_path
            done += 1
            progress.progress(min(1.0, done / total_steps))

        # Historia
        hist_entry = {
            "file": uploaded_pdf.name,
            "pages": [p + 1 for p in selected_pages],
            "tesseract_lang": tess_lang,
            "detected_lang": detected,
            "translated_to_pl": bool(do_translate),
            "exports": list(downloads.keys()),
            "timestamp": datetime.now().isoformat(timespec="seconds")
        }
        save_history(hist_entry)

        status.success("✅ Zakończono! Wyniki gotowe do pobrania poniżej.")
        progress.progress(1.0)

        st.subheader("💾 Pobieranie wyników")
        for label, path in downloads.items():
            st.write(f"{label}: {os.path.basename(path)}")
            st.download_button(
                label=f"💾 Pobierz {label}",
                data=bytes_from_file(path),
                file_name=os.path.basename(path),
                mime="application/octet-stream",
                key=f"dl_{label}"
            )

# Historia konwersji + czyszczenie
st.markdown("---")
st.subheader("📜 Historia konwersji")

col_h1, col_h2 = st.columns([1, 2])
with col_h1:
    confirm_clear = st.checkbox("Potwierdzam usunięcie historii")
with col_h2:
    if st.button("🗑️ Wyczyść historię"):
        if confirm_clear:
            try:
                clear_history()
                st.success("Historia została wyczyszczona.")
            except Exception as e:
                st.error(f"Nie udało się wyczyścić historii: {e}")
        else:
            st.warning("Zaznacz „Potwierdzam usunięcie historii”, aby kontynuować.")

history = load_history()
if history:
    for item in reversed(history[-10:]):
        with st.container(border=True):
            st.write(
                f"• Plik: {item['file']} • Strony: {item['pages']} "
                f"• OCR: {item['tesseract_lang']} • Wykryty: {item['detected_lang']} "
                f"• Tłum.: {item['translated_to_pl']} • Eksport: {item['exports']} "
                f"• {item['timestamp']}"
            )
else:
    st.caption("Brak historii. Wykonaj pierwszą konwersję.")

# Stopka
st.markdown("---")
st.markdown("Wygenerowano przez **PDF OCR Studio** — by `netdark_1966`")
